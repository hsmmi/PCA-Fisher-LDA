import os
import pandas as pd
import numpy as np
import docx
from io import StringIO
from dataset import Dataset
from normalization import clipping, log_scaling, range_min_to_max, z_score


def read_dataset(file, atr):
    if (type(atr) == int):
        with open(os.path.join(os.path.dirname(__file__), file), 'r') as f:
            return list(map(
                lambda x: float(x.split(', ')[atr]), f.read().splitlines()))
    else:
        with open(os.path.join(os.path.dirname(__file__), file), 'r') as f:
            return list(map(
                lambda x: [float(i) for i in (x.split(', ')[atr[0]:atr[1]])],
                f.read().splitlines()))


def dataframe_to_docx_table(header, data, file, doc=None, save=1):
    """
    Read header and data
    If you gave if doc it add header and data to it and return it
    If you gave it save=0 it will not be save doc
    Return doc include header and data
    """
    if(doc is None):
        doc = docx.Document()
    doc.add_heading(header, 1)

    table = doc.add_table(rows=len(data.index)+1, cols=len(data.columns)+1)

    for j in range(len(data.columns)):
        table.cell(0, j+1).text = f'{data.columns[j]}'

    for i in range(len(data.index)):
        table.cell(i+1, 0).text = f'{data.index[i]}'
        for j in range(len(data.columns)):
            table.cell(i+1, j+1).text = f'{data.iat[i, j]}'
    table.style = 'Table Grid'
    if(save):
        doc.save(file)
    return doc


def string_to_dataframe(string):
    data = StringIO(string)
    return pd.read_csv(data)


def nparray_to_csv(file: str, input: np.ndarray, decimal: int) -> None:
    path = os.path.join(os.path.dirname(__file__), file)
    pd.DataFrame(np.round(input, decimal)).to_csv(path)


def print_array_with_dataframe(array):
    print(pd.DataFrame(array))


def read_dataset_with_pandas(self, file, atr=None):
    colName = pd.read_csv(
        os.path.join(os.path.dirname(__file__), file), nrows=0).columns
    if (type(atr) == int):
        colName = [colName[atr]]
    elif(atr is not None):
        colName = colName[atr[0]:atr[1]]
    data = pd.read_csv(
        os.path.join(os.path.dirname(__file__), file), usecols=colName)

    return colName, data


def read_dataset_to_X_and_y(
        file, range_feature=None, range_label=None,
        normalization=None, min_value=None, max_value=None, add_x0=False,
        shuffle=False, about_nan=None):
    """
    Read the attribute(range_atr) that you want and put X0 = 1 and thoes
    attribute of all samples in X and all samples label in y
    normalization:
    .   by default is None and can be "z_score", "scaling", "clipping"
        or "log_scaling"
    .   for "scaling", "clipping" must set min_value and max_value
    Return X and y as nparray
    """
    data = pd.read_csv(os.path.join(os.path.dirname(__file__), file))
    col_name = data.columns

    dataset = Dataset()

    if(about_nan == 'delete'):
        data.dropna(inplace=True)

    for col in col_name:
        if(data[col].dtype != int and data[col].dtype != float):
            data[col] = pd.factorize(data[col])[0]

    data = data.to_numpy()

    if(shuffle is True):
        np.random.shuffle(data)

    number_of_attribute = len(col_name)
    if(range_feature is None):
        range_feature = (0, number_of_attribute-1)
    if(range_label is None):
        range_label = (number_of_attribute-1, number_of_attribute)

    dataset.feature_name = col_name[range_feature[0]:range_feature[1]]
    dataset.label_name = col_name[range_label[0]:range_label[1]]

    dataset.sample = np.array(list(map(
        lambda x: x[range_feature[0]:range_feature[1]], data)))
    dataset.label = np.array(list(map(
        lambda x: x[range_label[0]:range_label[1]], data)))

    if(about_nan == 'class_mean'):
        dataset.sample = dataset.sample.astype(float)
        for a_label in dataset.diffrent_label:
            class_label = dataset.sample[(dataset.label == a_label).flatten()]
            for a_feature in range(dataset.number_of_feature):
                mean_feature_label = np.nanmean(class_label[:, a_feature])
                for a_sample in range(dataset.number_of_sample):
                    if np.isnan(dataset.sample[a_sample, a_feature]):
                        dataset.sample[a_sample, a_feature] = \
                            mean_feature_label

    if(normalization is not None):
        if(normalization == 'z_score'):
            dataset.sample = z_score(dataset.sample)
        elif(normalization == 'scaling'):
            dataset.sample = range_min_to_max(
                dataset.sample, min_value, max_value)
        elif(normalization == 'clipping'):
            dataset.sample = clipping(dataset.sample, min_value, max_value)
        elif(normalization == 'logScaling'):
            dataset.sample = log_scaling(dataset.sample)
        else:
            print(
                'method should be "z_score", "scaling", "clipping" or'
                '"logScaling"')
            return

    dataset.diffrent_label = np.unique(dataset.label)
    dataset.number_of_feature = dataset.sample.shape[1]
    dataset.number_of_sample = dataset.sample.shape[0]

    if(add_x0 is True):
        dataset.sample = np.concatenate(
            (np.array([1.0]*dataset.number_of_sample), dataset.sample), axis=1)

    return dataset
